from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.text_processing import chunk_text, extract_text


def test_chunk_text_preserves_overlap() -> None:
    chunks = chunk_text("abcdefghijklmnopqrstuvwxyz", chunk_size=10, overlap=3)

    assert chunks == ["abcdefghij", "hijklmnopq", "opqrstuvwx", "vwxyz"]


def test_chunk_text_returns_empty_for_blank_text() -> None:
    assert chunk_text("   \n\t   ", chunk_size=10, overlap=2) == []


def test_chunk_text_rejects_invalid_overlap() -> None:
    with pytest.raises(ValueError, match="CHUNK_OVERLAP"):
        chunk_text("hello world", chunk_size=10, overlap=10)


def test_extract_text_from_txt(tmp_path: Path) -> None:
    file_path = tmp_path / "notes.txt"
    file_path.write_text("Content analysis notes", encoding="utf-8")

    assert extract_text(file_path) == "Content analysis notes"


def test_extract_text_from_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.save(file_path)

    assert extract_text(file_path) == "First paragraph\nSecond paragraph"


def test_extract_text_rejects_unsupported_file_type(tmp_path: Path) -> None:
    file_path = tmp_path / "image.png"
    file_path.write_bytes(b"not really an image")

    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(file_path)

